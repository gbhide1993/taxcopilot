from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import HTTPException
from sqlalchemy.orm import Session
from datetime import datetime
from io import BytesIO

from app.models.draft_version import DraftVersion
from app.models.notice import Notice
from app.models.client import Client
from app.models.sections_master import SectionsMaster
from app.models.firm import Firm
from app.models.appeal_versions import AppealVersion


def generate_draft_docx(db: Session, notice_id: int, version_number: int):

    # Fetch draft version
    draft = (
        db.query(DraftVersion)
        .filter(
            DraftVersion.notice_id == notice_id,
            DraftVersion.version_number == version_number
        )
        .first()
    )

    if not draft:
        raise HTTPException(status_code=404, detail="Draft version not found.")

    # Fetch related entities
    notice = db.query(Notice).filter(Notice.id == notice_id).first()
    client = db.query(Client).filter(Client.id == notice.client_id).first()
    firm = db.query(Firm).first()

    if not firm:
        raise HTTPException(status_code=500, detail="Firm configuration not found.")

    # Create document
    document = Document()

    # ==========================
    # Firm Header (Dynamic)
    # ==========================
    header = document.add_heading(firm.name, level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if firm.address:
        p = document.add_paragraph(firm.address)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if firm.email or firm.phone:
        contact_line = f"Email: {firm.email or ''}"
        if firm.phone:
            contact_line += f" | Phone: {firm.phone}"
        p = document.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # ==========================
    # Date
    # ==========================
    document.add_paragraph(
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}"
    )

    document.add_paragraph()

    # ==========================
    # Client Block
    # ==========================
    document.add_paragraph("To,")
    document.add_paragraph(client.name if client else "")
    document.add_paragraph()

    # ==========================
    # Subject
    # ==========================
    subject = document.add_paragraph()
    run = subject.add_run(
        f"Subject: Reply to Notice under Section {notice.section_reference}"
    )
    run.bold = True

    document.add_paragraph()

    # ==========================
    # Salutation
    # ==========================
    document.add_paragraph("Respected Sir/Madam,")
    document.add_paragraph()

    # ==========================
    # 1. Introduction
    # ==========================
    document.add_heading("1. Introduction", level=2)
    document.add_paragraph(draft.introduction)

    # ==========================
    # 2. Facts
    # ==========================
    document.add_heading("2. Facts of the Case", level=2)
    document.add_paragraph(draft.facts_summary)

    # ==========================
    # 3. Legal Position
    # ==========================
    document.add_heading("3. Legal Position", level=2)
    document.add_paragraph(draft.legal_position)

    # ==========================
    # 4. Relevant Provision
    # ==========================
    document.add_heading("4. Relevant Legal Provision", level=2)
    document.add_paragraph(draft.section_reference)

    # ==========================
    # 5. Prayer
    # ==========================
    document.add_heading("5. Prayer", level=2)
    document.add_paragraph(draft.prayer)

    document.add_paragraph()

    # ==========================
    # Annexure Auto Append
    # ==========================
    section = (
        db.query(SectionsMaster)
        .filter(
            SectionsMaster.act_name == notice.act_name,
            SectionsMaster.section_reference == notice.section_reference
        )
        .first()
    )

    if section and section.annexure_template:
        document.add_page_break()
        document.add_heading("Annexure", level=2)
        document.add_paragraph(section.annexure_template)

    # ==========================
    # Closing & Digital Signature
    # ==========================
    document.add_paragraph()
    document.add_paragraph("Thanking you.")
    document.add_paragraph()
    document.add_paragraph("Yours faithfully,")
    document.add_paragraph(f"For {firm.name}")
    document.add_paragraph()

    if firm.signature_name:
        document.add_paragraph(firm.signature_name)

    if firm.signature_designation:
        document.add_paragraph(firm.signature_designation)

    # ==========================
    # Footer Metadata
    # ==========================
    document.add_paragraph()
    document.add_paragraph(
        f"(Draft Version {version_number} | Generated on "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC)"
    )

    # Save to memory
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream

########## APPEAL DOCX GENERATION (Similar structure to draft but with appeal-specific content) ##########

def generate_appeal_docx(db: Session, notice_id: int, version_number: int):

    appeal = (
        db.query(AppealVersion)
        .filter(
            AppealVersion.notice_id == notice_id,
            AppealVersion.version_number == version_number
        )
        .first()
    )

    if not appeal:
        raise HTTPException(status_code=404, detail="Appeal version not found.")

    notice = db.query(Notice).filter(Notice.id == notice_id).first()
    client = db.query(Client).filter(Client.id == notice.client_id).first()
    firm = db.query(Firm).first()

    document = Document()

    # ==========================
    # Firm Header
    # ==========================
    header = document.add_heading(firm.name, level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if firm.address:
        p = document.add_paragraph(firm.address)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if firm.email or firm.phone:
        contact = f"Email: {firm.email or ''}"
        if firm.phone:
            contact += f" | Phone: {firm.phone}"
        p = document.add_paragraph(contact)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # ==========================
    # Title
    # ==========================
    title = document.add_heading("Grounds of Appeal", level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # ==========================
    # Appellant Details
    # ==========================
    document.add_paragraph(f"Appellant: {client.name}")
    document.add_paragraph(
        f"Against proceedings under Section {notice.section_reference} of {notice.act_name}"
    )

    document.add_paragraph()
    document.add_paragraph(f"Date: {datetime.utcnow().strftime('%d %B %Y')}")
    document.add_paragraph()

    # ==========================
    # Background
    # ==========================
    document.add_heading("Background", level=3)
    document.add_paragraph(appeal.background)

    document.add_paragraph()

    # ==========================
    # Grounds
    # ==========================
    document.add_heading("Grounds of Appeal", level=3)
    document.add_paragraph(appeal.grounds)

    document.add_paragraph()

    # ==========================
    # Prayer
    # ==========================
    document.add_heading("Prayer", level=3)
    document.add_paragraph(appeal.prayer)

    document.add_paragraph()

    # ==========================
    # Signature
    # ==========================
    document.add_paragraph("Yours faithfully,")
    document.add_paragraph(f"For {firm.name}")
    document.add_paragraph()

    if firm.signature_name:
        document.add_paragraph(firm.signature_name)

    if firm.signature_designation:
        document.add_paragraph(firm.signature_designation)

    document.add_paragraph()
    document.add_paragraph(
        f"(Appeal Version {version_number} | Generated on "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC)"
    )

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream
